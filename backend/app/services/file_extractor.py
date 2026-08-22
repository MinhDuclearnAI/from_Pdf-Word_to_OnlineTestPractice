import os
import logging
import fitz  # PyMuPDF
import docx
from typing import Tuple
from io import BytesIO

logger = logging.getLogger(__name__)

class FileParsingError(Exception):
    """Custom exception khi không thể đọc hoặc parse file."""
    pass

BOLD_MARKER_START = "[[BOLD_START]]"
BOLD_MARKER_END = "[[BOLD_END]]"

def strip_structure_markers(text: str) -> str:
    """Loại bỏ các marker cấu trúc khỏi text trước khi lưu hoặc hiển thị."""
    if not text:
        return ""
    text = text.replace(BOLD_MARKER_START, "")
    text = text.replace(BOLD_MARKER_END, "")
    return text

def clean_and_normalize_text(raw_text: str) -> str:
    """
    Chuẩn hóa khoảng trắng ngang và nối liền các dòng bị ngắt mềm (soft-linebreaks) do lề PDF,
    kể cả Danh từ riêng viết hoa, Số liệu, Từ ghép có gạch nối; 
    đồng thời bảo toàn 100% các dòng trống tách đoạn (\n\n) và cấu trúc tiêu đề, câu hỏi.
    """
    if not raw_text or not raw_text.strip():
        return ""

    import re
    # 1. Xóa khoảng trắng ngang thừa (spaces, tabs) nhưng giữ lại \n
    text = re.sub(r'[^\S\n]+', ' ', raw_text)
    
    # 2. Xử lý thông minh các dòng bị ngắt mềm do lề PDF
    lines = text.split('\n')
    unwrapped = []
    
    for line in lines:
        line_str = line.strip()
        if not line_str:
            unwrapped.append("") # Bảo toàn tách đoạn \n\n
            continue
            
        if not unwrapped or unwrapped[-1] == "":
            unwrapped.append(line_str)
            continue
            
        prev = unwrapped[-1]
        
        # Kiểm tra nếu dòng hiện tại là bắt đầu của một khối mới (câu hỏi, đáp án [A-Z], tiêu đề) -> KHÔNG nối
        is_new_block_start = bool(re.match(
            r'^(Questions?\s+\d+|\d+[\.\)]|[A-Z][\.\)]|[A-Z]\s+|READING\s+PASSAGE|TRUE|FALSE|NOT GIVEN|YES|NO|Choose|Complete|Write|List\s+of)', 
            line_str, 
            re.I
        ))
        
        # Kiểm tra nếu dòng trước là tiêu đề / header đơn lẻ (< 80 ký tự)
        is_prev_header = bool(re.match(r'^(READING\s+PASSAGE|Questions?\s+\d+)', prev, re.I)) or (
            len(prev) < 80 and (prev.endswith('!') or not prev.endswith(('.', ':', ';', ',')) and len(prev.split()) <= 8)
        )
        
        if is_new_block_start or is_prev_header:
            unwrapped.append(line_str)
            continue
            
        # Nếu dòng trước kết thúc bằng dấu gạch nối (từ bị bẻ đôi qua lề: e.g. "Triassic-", "disease-")
        if prev.endswith('-'):
            unwrapped[-1] = prev + line_str
        elif not prev.endswith(('.', '!', '?', ':')):
            # Dòng trước chưa kết thúc câu (dù kết thúc bằng từ thường, từ viết hoa hay số liệu) -> Nối với 1 khoảng trắng
            unwrapped[-1] = prev + " " + line_str
        else:
            # Dòng trước kết thúc bằng dấu chấm nhưng trong cùng một đoạn văn (không có dòng trống \n\n)
            unwrapped[-1] = prev + " " + line_str
            
    result = "\n".join(unwrapped)
    # Rút gọn các dòng trống liên tiếp (3 trở lên) thành 2 dòng trống (chuẩn tách đoạn \n\n)
    result = re.sub(r'\n{3,}', '\n\n', result)
    return result.strip()

def is_continuation(prev_text: str, next_text: str) -> bool:
    if not prev_text or not next_text:
        return False
    prev_text = prev_text.strip()
    next_text = next_text.strip()
    if not prev_text or not next_text:
        return False

    import re
    # 1. BẢO VỆ TUYỆT ĐỐI: Nếu khối sau là bắt đầu của 1 câu hỏi, đáp án, header -> KHÔNG BAO GIỜ GỘP!
    if re.match(r'^(Questions?\s+\d+|\d+[\.\)]|[A-Z][\.\)]|[A-Z]\s+|READING\s+PASSAGE|TRUE|FALSE|NOT GIVEN|YES|NO|Choose|Complete|Write|List\s+of)', next_text, re.I):
        return False

    # 2. BẢO VỆ TIÊU ĐỀ: Nếu dòng trước là tiêu đề đứng độc lập (< 80 ký tự)
    if len(prev_text) < 80 and (prev_text.endswith('!') or not prev_text.endswith(('.', ':', ';', ',')) and len(prev_text.split()) <= 8):
        return False

    # 3. Nối từ bị bẻ đôi bằng gạch nối
    if prev_text.endswith("-"):
        return True
        
    if next_text[0].islower():
        return True
        
    last_word = prev_text.split()[-1].lower() if prev_text.split() else ""
    if last_word in ["and", "or", "but", "the", "a", "an", "of", "in", "to", "with", "for", "on", "at", "by", "from", "as", "is", "are"]:
        return True
        
    if prev_text.endswith(","):
        return True

    return False

def extract_text_from_pdf(file_path: str) -> Tuple[str, bool]:
    """
    Trích xuất text từ file PDF sử dụng PyMuPDF với khả năng nhận diện bố cục 2 cột (Multi-column Sorter).
    Trả về Tuple[text_đã_trích_xuất, is_scanned_pdf].
    """
    text_content = []
    is_scanned = True  # Giả định ban đầu là file scan

    try:
        from app.services.storage_service import storage
        
        # 1. Chạy pipeline inventory để tìm và crop ảnh chuẩn xác nhất (kể cả ảnh lẩn trong text)
        inventory = inventory_page(file_path)
        assets = extract_all_images_from_inventory(file_path, inventory)
        
        # Nhóm assets theo page
        page_assets = {}
        for asset in assets:
            p = asset["page"]
            if p not in page_assets:
                page_assets[p] = []
            page_assets[p].append(asset)

        doc = fitz.open(file_path)
        import re
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            page_width = page.rect.width
            page_height = page.rect.height
            
            # Lấy danh sách blocks thô từ PyMuPDF
            blocks = page.get_text("blocks", sort=False)
            
            images_in_page = page_assets.get(page_num, [])
            
            # 2. Thu thập các text blocks và image items trên trang
            page_items = []
            
            for b in blocks:
                x0, y0, x1, y1, text_data, block_no, block_type = b
                if block_type == 0: # Text block
                    page_text = text_data.strip()
                    if not page_text:
                        continue
                        
                    # Lọc rác Header/Footer (Page number, URL)
                    is_in_margin = y0 < 50 or y1 > page_height - 50
                    if is_in_margin and len(page_text) < 80:
                        pt_lower = page_text.lower()
                        is_page_num = bool(re.match(r'^\d+$', page_text)) or "page" in pt_lower
                        is_url = bool(re.match(r'^https?://', pt_lower)) or bool(re.match(r'^www\.', pt_lower)) or ".com" in pt_lower or ".org" in pt_lower or ".vn" in pt_lower
                        if is_page_num or is_url:
                            continue
                            
                    page_items.append({
                        "type": "text",
                        "x0": x0,
                        "y0": y0,
                        "x1": x1,
                        "y1": y1,
                        "content": page_text
                    })
                    is_scanned = False
            
            for asset in images_in_page:
                page_items.append({
                    "type": "image",
                    "x0": asset.get("x0", 0),
                    "y0": asset["y0"],
                    "x1": asset.get("x1", page_width),
                    "y1": asset.get("y1", asset["y0"] + 100),
                    "url": asset["image_url"]
                })
                
            # 3. THUẬT TOÁN PHÁT HIỆN & SẮP XẾP BỐ CỤC (SINGLE COLUMN vs TWO COLUMNS)
            left_col = [item for item in page_items if item["x1"] <= page_width * 0.55 and item["x0"] < page_width * 0.45]
            right_col = [item for item in page_items if item["x0"] >= page_width * 0.45 and item["x1"] > page_width * 0.55]
            
            is_two_column = len(left_col) >= 2 and len(right_col) >= 2
            
            if is_two_column:
                # Tìm Y nơi bắt đầu phân 2 cột
                col_top_y = min(min(item["y0"] for item in left_col), min(item["y0"] for item in right_col))
                
                # Header items (nằm trên vùng 2 cột hoặc trải rộng ngang trang)
                header_items = [
                    item for item in page_items 
                    if item["y1"] <= col_top_y + 10 or (item["x0"] < page_width * 0.35 and item["x1"] > page_width * 0.65)
                ]
                header_ids = set(id(x) for x in header_items)
                
                # Cột Trái và Cột Phải (loại bỏ header)
                c_left = [item for item in page_items if id(item) not in header_ids and item["x0"] < page_width * 0.5]
                c_right = [item for item in page_items if id(item) not in header_ids and item["x0"] >= page_width * 0.5]
                
                header_items.sort(key=lambda item: item["y0"])
                c_left.sort(key=lambda item: item["y0"])
                c_right.sort(key=lambda item: item["y0"])
                
                ordered_items = header_items + c_left + c_right
            else:
                # Trang 1 cột chuẩn: sắp xếp tuần tự theo chiều dọc Y từ trên xuống dưới
                ordered_items = sorted(page_items, key=lambda item: item["y0"])
            
            for item in ordered_items:
                if item["type"] == "image":
                    text_content.append(f"\n[[IMAGE_REF: {item['url']}]]\n")
                else:
                    page_text = item["content"]
                    if not text_content:
                        text_content.append(page_text)
                    else:
                        prev_text = text_content[-1]
                        if "[[IMAGE_REF" not in prev_text and is_continuation(prev_text, page_text):
                            if text_content[-1].endswith("-"):
                                text_content[-1] = text_content[-1][:-1] + page_text
                            else:
                                text_content[-1] = text_content[-1] + " " + page_text
                        else:
                            text_content.append(page_text)

        doc.close()
        
        raw_text = "\n\n".join(text_content)
        normalized_text = clean_and_normalize_text(raw_text)
        
        return normalized_text, is_scanned

    except Exception as e:
        logger.error(f"Lỗi khi đọc file PDF {file_path}: {str(e)}")
        raise FileParsingError(f"Không thể trích xuất PDF: {str(e)}")

def inventory_page(file_path: str) -> list:
    """
    Kiểm kê từng trang PDF: có ảnh nhúng không, có vùng vẽ vector không, độ dài text.
    """
    inventory = []
    try:
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            images = page.get_images()
            drawings = page.get_drawings()
            text = page.get_text("text")
            
            image_rects = []
            for img_info in images:
                xref = img_info[0]
                rects = page.get_image_rects(xref)
                for rect in rects:
                    image_rects.append({
                        "x0": rect.x0, "y0": rect.y0, "x1": rect.x1, "y1": rect.y1
                    })
            
            drawing_rects = []
            for d in drawings:
                rect = d["rect"]
                drawing_rects.append({
                    "x0": rect.x0, "y0": rect.y0, "x1": rect.x1, "y1": rect.y1
                })
            
            inventory.append({
                "page": page_num,
                "has_image": len(images) > 0,
                "has_vector": len(drawings) > 0,
                "text_len": len(text.strip()),
                "image_rects": image_rects,
                "drawing_rects": [] # Bỏ qua vector drawings vì PyMuPDF lấy cả các đường kẻ viền bảng, sinh ra rác
            })
        doc.close()
    except Exception as e:
        logger.error(f"Lỗi khi kiểm kê PDF {file_path}: {e}")
    return inventory

def process_image_blocks(file_path: str, image_blocks: list) -> list:
    """
    (Deprecated) Dùng extract_all_images_from_inventory thay thế.
    """
    pass

def extract_all_images_from_inventory(file_path: str, inventory: list) -> list:
    """
    Tự động crop tất cả ảnh từ inventory (tránh mất ảnh), upload và trả về list URL.
    """
    assets_found = []
    try:
        from app.services.storage_service import storage
        doc = fitz.open(file_path)
        
        img_counter = 1
        for page_data in inventory:
            page_num = page_data.get("page", 0)
            rects = page_data.get("image_rects", []) + page_data.get("drawing_rects", [])
            
            if not rects:
                continue
                
            page = doc.load_page(page_num)
            for rect_dict in rects:
                rect = fitz.Rect(rect_dict["x0"], rect_dict["y0"], rect_dict["x1"], rect_dict["y1"])
                try:
                    cropped_pix = page.get_pixmap(clip=rect, dpi=150)
                    cropped_bytes = cropped_pix.tobytes("jpeg")
                    
                    file_obj = BytesIO(cropped_bytes)
                    original_name = f"auto_crop_p{page_num}_{img_counter}.jpg"
                    object_name = storage.upload_file(file_obj, original_name, content_type="image/jpeg", folder="exams/diagrams")
                    
                    image_url = storage.get_presigned_url(object_name, expiration=7*24*3600)
                    
                    if image_url:
                        assets_found.append({
                            "page": page_num,
                            "image_url": image_url,
                            "y0": rect_dict["y0"],
                            "x0": rect_dict["x0"]
                        })
                    img_counter += 1
                except Exception as crop_err:
                    logger.error(f"Lỗi khi crop ảnh tại trang {page_num}: {crop_err}")
                    
        doc.close()
    except Exception as e:
        logger.error(f"Lỗi khi trích xuất toàn bộ ảnh từ {file_path}: {e}")
        
    return assets_found

def extract_text_from_docx(file_path: str) -> str:
    """
    Trích xuất text từ file Word (.docx), bao gồm cả các đoạn văn thông thường
    và dữ liệu nằm trong các bảng biểu (Tables).
    """
    text_content = []

    try:
        doc = docx.Document(file_path)
        
        # 1. Đọc các đoạn văn bản (Paragraphs)
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)

        # 2. Đọc dữ liệu trong các bảng biểu (rất phổ biến trong đề thi Tiếng Anh/Văn)
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    if cell.text.strip():
                        # Thay thế \n trong cell thành khoảng trắng để không làm vỡ format
                        row_data.append(cell.text.strip().replace("\n", " "))
                if row_data:
                    text_content.append(" | ".join(row_data))
                    
        raw_text = "\n".join(text_content)
        return clean_and_normalize_text(raw_text)

    except Exception as e:
        logger.error(f"Lỗi khi đọc file DOCX {file_path}: {str(e)}")
        raise FileParsingError(f"Không thể trích xuất DOCX: {str(e)}")

def process_file(file_path: str, mime_type: str) -> str:
    """
    Hàm main entry point để gọi từ Celery Worker.
    Dựa vào MIME type để điều hướng sang hàm trích xuất tương ứng.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File không tồn tại trên hệ thống: {file_path}")

    logger.info(f"Bắt đầu trích xuất file: {file_path} (MIME: {mime_type})")

    if mime_type == "application/pdf":
        extracted_text, is_scanned = extract_text_from_pdf(file_path)
        
        if is_scanned:
            logger.warning(f"File {file_path} là định dạng PDF Scan/Ảnh.")
            # TODO: Implement trigger OCR flow tại đây (ví dụ gọi ocr_service.py)
            # return ocr_service.process_scanned_pdf(file_path)
            
        return extracted_text

    elif mime_type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
        "application/msword"
    ]:
        return extract_text_from_docx(file_path)
    
    else:
        raise FileParsingError(f"Định dạng file không được hỗ trợ: {mime_type}")